from pathlib import Path
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from backend.config import settings


def parse_file(file_path: str) -> List[Document]:
    """解析文件，返回 LangChain Document 列表"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix == ".docx":
        return _parse_docx(file_path)
    elif suffix == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {suffix}")


def _parse_pdf(file_path: str) -> List[Document]:
    import fitz  # PyMuPDF
    docs = []
    pdf = fitz.open(file_path)
    for page_num, page in enumerate(pdf):
        text = page.get_text("text").strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "page": page_num + 1}
            ))
    pdf.close()
    return docs


def _parse_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=full_text, metadata={"source": file_path, "page": 1})]


def _parse_txt(file_path: str) -> List[Document]:
    text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source": file_path, "page": 1})]


def split_documents(docs: List[Document]) -> List[Document]:
    """将文档切分为小块"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
    )
    return splitter.split_documents(docs)
